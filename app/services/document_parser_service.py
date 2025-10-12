"""
Document parser service for handling various file formats.
Supports PDF, DOCX, TXT, CSV, XLSX with intelligent content extraction.
"""

import logging
import io
import chardet
from typing import Dict, Any, List, Optional, Tuple
from pathlib import Path
try:
    import magic
except ImportError:
    # Fallback for Windows - use python-magic-bin
    try:
        import magic
    except ImportError:
        magic = None

# Document processing imports
import PyPDF2
from docx import Document
import pandas as pd
import openpyxl

logger = logging.getLogger(__name__)


class DocumentParserService:
    """Service for parsing various document formats."""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self._parse_pdf,
            'docx': self._parse_docx,
            'txt': self._parse_txt,
            'csv': self._parse_csv,
            'xlsx': self._parse_xlsx,
            'xls': self._parse_xlsx
        }
    
    async def parse_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse a document and extract content with metadata.
        
        Args:
            file_content: Raw file content as bytes
            filename: Original filename
            
        Returns:
            Dictionary with parsed content and metadata
        """
        try:
            # Detect file type
            file_type = self._detect_file_type(file_content, filename)
            
            if file_type not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_type}")
            
            # Parse the document
            parser_func = self.supported_formats[file_type]
            parsed_content = await parser_func(file_content, filename)
            
            # Add common metadata
            parsed_content.update({
                "file_type": file_type,
                "filename": filename,
                "file_size": len(file_content),
                "parsing_method": f"document_parser_{file_type}"
            })
            
            logger.info(f"Successfully parsed {file_type} document: {filename}")
            return parsed_content
            
        except Exception as e:
            logger.error(f"Error parsing document {filename}: {e}")
            raise
    
    def _detect_file_type(self, file_content: bytes, filename: str) -> str:
        """Detect file type using magic bytes and extension."""
        try:
            # Try magic bytes first if available
            if magic is not None:
                mime_type = magic.from_buffer(file_content, mime=True)
                
                # Map MIME types to our supported formats
                mime_mapping = {
                    'application/pdf': 'pdf',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
                    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': 'xlsx',
                    'application/vnd.ms-excel': 'xls',
                    'text/plain': 'txt',
                    'text/csv': 'csv',
                    'application/csv': 'csv'
                }
                
                if mime_type in mime_mapping:
                    return mime_mapping[mime_type]
            
            # Fallback to file extension
            extension = Path(filename).suffix.lower().lstrip('.')
            if extension in self.supported_formats:
                return extension
            
            raise ValueError(f"Could not detect file type for {filename}")
            
        except Exception as e:
            logger.error(f"Error detecting file type: {e}")
            raise
    
    async def _parse_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse PDF document."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract text from all pages
            full_text = ""
            page_texts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                page_texts.append({
                    "page_number": page_num + 1,
                    "content": page_text,
                    "char_count": len(page_text)
                })
                full_text += page_text + "\n\n"
            
            # Extract metadata
            metadata = {}
            if pdf_reader.metadata:
                metadata = {
                    "title": pdf_reader.metadata.get("/Title", ""),
                    "author": pdf_reader.metadata.get("/Author", ""),
                    "subject": pdf_reader.metadata.get("/Subject", ""),
                    "creator": pdf_reader.metadata.get("/Creator", ""),
                    "producer": pdf_reader.metadata.get("/Producer", ""),
                    "creation_date": str(pdf_reader.metadata.get("/CreationDate", "")),
                    "modification_date": str(pdf_reader.metadata.get("/ModDate", ""))
                }
            
            return {
                "content": full_text.strip(),
                "title": metadata.get("title") or Path(filename).stem,
                "metadata": metadata,
                "structure": {
                    "total_pages": len(pdf_reader.pages),
                    "page_contents": page_texts
                },
                "source": filename
            }
            
        except Exception as e:
            logger.error(f"Error parsing PDF {filename}: {e}")
            raise
    
    async def _parse_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse DOCX document."""
        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            # Extract text content
            full_text = ""
            paragraphs = []
            tables = []
            
            # Process paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name if para.style else "Normal",
                        "char_count": len(para.text)
                    })
                    full_text += para.text + "\n"
            
            # Process tables
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                tables.append({
                    "table_index": table_idx,
                    "data": table_data,
                    "rows": len(table_data),
                    "columns": len(table_data[0]) if table_data else 0
                })
                
                # Add table content to full text
                for row in table_data:
                    full_text += " | ".join(row) + "\n"
            
            # Extract metadata
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or Path(filename).stem,
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "comments": core_props.comments or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else ""
            }
            
            return {
                "content": full_text.strip(),
                "title": metadata["title"],
                "metadata": metadata,
                "structure": {
                    "paragraphs": paragraphs,
                    "tables": tables,
                    "total_paragraphs": len(paragraphs),
                    "total_tables": len(tables)
                },
                "source": filename
            }
            
        except Exception as e:
            logger.error(f"Error parsing DOCX {filename}: {e}")
            raise
    
    async def _parse_txt(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse TXT document."""
        try:
            # Detect encoding
            encoding_result = chardet.detect(file_content)
            encoding = encoding_result.get('encoding', 'utf-8')
            confidence = encoding_result.get('confidence', 0)
            
            # Try to decode with detected encoding
            try:
                text_content = file_content.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                # Fallback to utf-8 with error handling
                text_content = file_content.decode('utf-8', errors='replace')
                encoding = 'utf-8'
                confidence = 0.5
            
            # Split into lines for structure analysis
            lines = text_content.split('\n')
            paragraphs = []
            current_paragraph = ""
            
            for line in lines:
                line = line.strip()
                if line:
                    current_paragraph += line + " "
                else:
                    if current_paragraph:
                        paragraphs.append({
                            "text": current_paragraph.strip(),
                            "char_count": len(current_paragraph.strip())
                        })
                        current_paragraph = ""
            
            # Add last paragraph if exists
            if current_paragraph:
                paragraphs.append({
                    "text": current_paragraph.strip(),
                    "char_count": len(current_paragraph.strip())
                })
            
            return {
                "content": text_content.strip(),
                "title": Path(filename).stem,
                "metadata": {
                    "encoding": encoding,
                    "encoding_confidence": confidence,
                    "line_count": len(lines),
                    "paragraph_count": len(paragraphs)
                },
                "structure": {
                    "paragraphs": paragraphs,
                    "total_paragraphs": len(paragraphs)
                },
                "source": filename
            }
            
        except Exception as e:
            logger.error(f"Error parsing TXT {filename}: {e}")
            raise
    
    async def _parse_csv(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse CSV document."""
        try:
            # Detect encoding
            encoding_result = chardet.detect(file_content)
            encoding = encoding_result.get('encoding', 'utf-8')
            
            # Try to decode and parse CSV
            try:
                text_content = file_content.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                text_content = file_content.decode('utf-8', errors='replace')
                encoding = 'utf-8'
            
            # Parse CSV with pandas
            csv_io = io.StringIO(text_content)
            
            # Try different separators
            separators = [',', ';', '\t', '|']
            df = None
            
            for sep in separators:
                try:
                    csv_io.seek(0)
                    df = pd.read_csv(csv_io, sep=sep, encoding=encoding)
                    if len(df.columns) > 1:  # Valid CSV with multiple columns
                        break
                except Exception:
                    continue
            
            if df is None or len(df.columns) <= 1:
                # Fallback: treat as plain text
                return await self._parse_txt(file_content, filename)
            
            # Extract structured data
            columns = df.columns.tolist()
            rows = df.values.tolist()
            
            # Create readable text representation
            text_content = ""
            text_content += " | ".join(columns) + "\n"
            text_content += "-" * (len(" | ".join(columns))) + "\n"
            
            for row in rows:
                text_content += " | ".join([str(cell) for cell in row]) + "\n"
            
            # Extract metadata
            metadata = {
                "encoding": encoding,
                "separator": sep,
                "total_rows": len(df),
                "total_columns": len(columns),
                "column_names": columns,
                "data_types": df.dtypes.astype(str).to_dict()
            }
            
            return {
                "content": text_content.strip(),
                "title": Path(filename).stem,
                "metadata": metadata,
                "structure": {
                    "columns": columns,
                    "rows": rows,
                    "total_rows": len(df),
                    "total_columns": len(columns)
                },
                "source": filename
            }
            
        except Exception as e:
            logger.error(f"Error parsing CSV {filename}: {e}")
            raise
    
    async def _parse_xlsx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Parse XLSX/XLS document."""
        try:
            # Load workbook
            workbook = openpyxl.load_workbook(io.BytesIO(file_content), data_only=True)
            
            all_sheets_content = ""
            sheets_data = []
            
            # Process each sheet
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # Extract data from sheet
                sheet_data = []
                max_row = sheet.max_row
                max_col = sheet.max_column
                
                for row in range(1, max_row + 1):
                    row_data = []
                    for col in range(1, max_col + 1):
                        cell_value = sheet.cell(row=row, column=col).value
                        row_data.append(str(cell_value) if cell_value is not None else "")
                    sheet_data.append(row_data)
                
                # Create readable text for this sheet
                sheet_text = f"Sheet: {sheet_name}\n"
                if sheet_data:
                    # Use first row as headers if it looks like headers
                    headers = sheet_data[0] if sheet_data else []
                    data_rows = sheet_data[1:] if len(sheet_data) > 1 else []
                    
                    if headers:
                        sheet_text += " | ".join(headers) + "\n"
                        sheet_text += "-" * (len(" | ".join(headers))) + "\n"
                    
                    for row in data_rows:
                        sheet_text += " | ".join(row) + "\n"
                
                sheets_data.append({
                    "sheet_name": sheet_name,
                    "data": sheet_data,
                    "max_row": max_row,
                    "max_col": max_col,
                    "content": sheet_text
                })
                
                all_sheets_content += sheet_text + "\n\n"
            
            # Extract metadata
            metadata = {
                "total_sheets": len(workbook.sheetnames),
                "sheet_names": workbook.sheetnames,
                "properties": {
                    "creator": workbook.properties.creator or "",
                    "title": workbook.properties.title or "",
                    "subject": workbook.properties.subject or "",
                    "description": workbook.properties.description or "",
                    "keywords": workbook.properties.keywords or "",
                    "created": str(workbook.properties.created) if workbook.properties.created else "",
                    "modified": str(workbook.properties.modified) if workbook.properties.modified else ""
                }
            }
            
            return {
                "content": all_sheets_content.strip(),
                "title": metadata["properties"]["title"] or Path(filename).stem,
                "metadata": metadata,
                "structure": {
                    "sheets": sheets_data,
                    "total_sheets": len(workbook.sheetnames)
                },
                "source": filename
            }
            
        except Exception as e:
            logger.error(f"Error parsing XLSX {filename}: {e}")
            raise
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        return list(self.supported_formats.keys())
    
    def validate_file_size(self, file_content: bytes, max_size_mb: int = 50) -> bool:
        """Validate file size."""
        max_size_bytes = max_size_mb * 1024 * 1024
        return len(file_content) <= max_size_bytes


# Global instance
document_parser_service = DocumentParserService()
